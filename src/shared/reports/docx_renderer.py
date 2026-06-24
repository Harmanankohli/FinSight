# ruff: noqa: E402
"""DOCX report generation."""

from __future__ import annotations

import logging
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)

from shared.reports.deck_model import DeckData
from shared.reports.deck_model import Section as DeckSection
from shared.reports.extraction import _extract_deck_data

# ── Colour palette (kept for backward compat) ─────────────────────────────────
_NAVY = "0C1425"
_BLUE = "3B7DFA"
_GREEN_DARK = "15803D"
_RED = "EF4444"
_TEXT = "1A2236"
_TEXT_LIGHT = "94A3B8"
_SURFACE = "F1F5F9"
_WHITE = "FFFFFF"

# DOCX uses a slightly different palette (kept for backward compat)
_DOCX_NAVY = "0C1425"
_DOCX_BLUE = "3B7DFA"
_DOCX_SURFACE = "F1F5F9"

_SIGNAL_COLOURS = {
    "BUY": (_GREEN_DARK, "DCFCE7"),
    "HOLD": ("1D4ED8", "DBEAFE"),
    "SELL": (_RED, "FEF2F2"),
}


def _extract_docx_content(
    brief_data: dict[str, Any],
    ticker: str,
    recommendation: str,
    confidence: float,
    analysis_date: str,
    company_info: dict[str, Any] | None = None,
) -> DeckData:
    """Wrapper for DOCX — reuses the deck extractor."""
    return _extract_deck_data(
        brief_data, ticker, recommendation, confidence, analysis_date, company_info=company_info
    )


def generate_docx(
    brief_data: dict[str, Any],
    ticker: str,
    recommendation: str,
    confidence: float,
    analysis_date: str,
    company_info: dict[str, Any] | None = None,
) -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    def rgb(hex_str: str) -> RGBColor:
        r = int(hex_str[0:2], 16)
        g = int(hex_str[2:4], 16)
        b = int(hex_str[4:6], 16)
        return RGBColor(r, g, b)

    logger.info("Generating DOCX report for %s", ticker)
    deck = _extract_docx_content(
        brief_data, ticker, recommendation, confidence, analysis_date, company_info=company_info
    )
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def set_run_style(
        run: Any,
        size: int,
        bold: bool = False,
        color: str = _TEXT,
        italic: bool = False,
        name: str = "Calibri",
    ) -> None:
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = rgb(color)

    # ── Document title block ──────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Investment Analysis Report")
    set_run_style(run, 28, bold=True, color=_DOCX_NAVY)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(ticker)
    set_run_style(run, 22, bold=True, color=_DOCX_BLUE)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(f"{analysis_date}  ·  Confidence: {confidence:.0%}")
    set_run_style(run, 11, color=_TEXT_LIGHT)

    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fg, _ = _SIGNAL_COLOURS.get(deck.recommendation, (_TEXT, _SURFACE))
    run = sig_p.add_run(f"  {deck.recommendation}  ")
    set_run_style(run, 14, bold=True, color=fg)

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    if deck.executive_summary:
        h = doc.add_heading("Executive Summary", level=1)
        h.runs[0].font.color.rgb = rgb(_DOCX_NAVY)
        h.runs[0].font.name = "Calibri"

        para = doc.add_paragraph(deck.executive_summary)
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "Calibri"
        para.runs[0].font.color.rgb = rgb(_TEXT)
        para.paragraph_format.space_after = Pt(8)
        doc.add_paragraph()

    # ── Key Metrics table ─────────────────────────────────────────────────────
    if deck.kpi_chips:
        h = doc.add_heading("Key Metrics", level=1)
        h.runs[0].font.color.rgb = rgb(_DOCX_NAVY)
        h.runs[0].font.name = "Calibri"

        items = [(c["label"], c["value"]) for c in deck.kpi_chips]
        tbl = doc.add_table(rows=len(items) + 1, cols=2)
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for ci, txt in enumerate(["Metric", "Value"]):
            hdr_cells[ci].text = txt
            run = hdr_cells[ci].paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = rgb(_WHITE)
            tc = hdr_cells[ci]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), _DOCX_BLUE)
            tcPr.append(shd)

        for ri, (label, value) in enumerate(items):
            row_cells = tbl.rows[ri + 1].cells
            row_cells[0].text = label
            row_cells[1].text = value
            for cell in row_cells:
                r = cell.paragraphs[0].runs[0]
                r.font.size = Pt(10)
                r.font.name = "Calibri"
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[0].paragraphs[0].runs[0].font.color.rgb = rgb(_DOCX_NAVY)
            if ri % 2 == 0:
                for cell in row_cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), _DOCX_SURFACE)
                    tcPr.append(shd)

        doc.add_paragraph()

    # ── Analysis sections ─────────────────────────────────────────────────────
    deck_section: DeckSection
    for deck_section in deck.sections:
        if not deck_section.title and not deck_section.body:
            continue
        if deck_section.title:
            h = doc.add_heading(deck_section.title, level=2)
            for run in h.runs:
                run.font.color.rgb = rgb(_DOCX_BLUE)
                run.font.name = "Calibri"
        if deck_section.body:
            for block in deck_section.body.split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("•"):
                    for line in block.split("\n"):
                        line = line.strip().lstrip("•").strip()
                        if line:
                            para = doc.add_paragraph(line, style="List Bullet")
                            if para.runs:
                                para.runs[0].font.size = Pt(11)
                                para.runs[0].font.name = "Calibri"
                else:
                    para = doc.add_paragraph(block)
                    if para.runs:
                        para.runs[0].font.size = Pt(11)
                        para.runs[0].font.name = "Calibri"
                        para.runs[0].font.color.rgb = rgb(_TEXT)
                    para.paragraph_format.space_after = Pt(6)
        doc.add_paragraph()

    # ── Risk / Opportunities ──────────────────────────────────────────────────
    if deck.opportunities or deck.risks:
        if deck.opportunities:
            h = doc.add_heading("Growth Opportunities", level=2)
            for run in h.runs:
                run.font.color.rgb = rgb(_GREEN_DARK)
                run.font.name = "Calibri"
            for opp in deck.opportunities:
                para = doc.add_paragraph(opp, style="List Bullet")
                if para.runs:
                    para.runs[0].font.size = Pt(11)
                    para.runs[0].font.name = "Calibri"

        if deck.risks:
            h = doc.add_heading("Key Risks", level=2)
            for run in h.runs:
                run.font.color.rgb = rgb(_RED)
                run.font.name = "Calibri"
            for risk in deck.risks:
                para = doc.add_paragraph(risk, style="List Bullet")
                if para.runs:
                    para.runs[0].font.size = Pt(11)
                    para.runs[0].font.name = "Calibri"

        doc.add_paragraph()

    # ── Disclaimer ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    h = doc.add_heading("Disclaimer", level=2)
    for run in h.runs:
        run.font.color.rgb = rgb(_DOCX_BLUE)
        run.font.name = "Calibri"
    para = doc.add_paragraph(deck.disclaimer)
    if para.runs:
        para.runs[0].font.size = Pt(9)
        para.runs[0].font.italic = True
        para.runs[0].font.color.rgb = rgb(_TEXT_LIGHT)
        para.runs[0].font.name = "Calibri"

    section_obj = doc.sections[0]
    footer = section_obj.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(f"Generated by FinSight  ·  {analysis_date}")
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(_TEXT_LIGHT)
    run.font.name = "Calibri"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("DOCX report generated for %s (%d bytes)", ticker, buf.getbuffer().nbytes)
    return buf
